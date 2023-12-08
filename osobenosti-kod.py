#---------подчищаем лишние------>(всё что до первой гл1)-----------
with open('1.txt', 'r', encoding=x) as f:
    lines = f.readlines()
index = -1
for i, line in enumerate(lines):
    if "1.1" in line:
        index = i
        break
if index > 1:
    del lines[:index-2]
with open('1.txt', 'w') as f:
    f.writelines(lines)





#--------------ЗАПИСЬ ГЛ И ПОДГЛ В ПЕРЕМЕННЫЕ--------------
# Открываем файл для чтения
with open('1.txt', 'r', encoding=x) as file:
    # Инициализируем переменные
    current_variable = ""
    ww_vars = []
    # Читаем файл построчно
    for line in file:
        # Проверяем, содержит ли строка цифру
        if any(char.isdigit() for char in line):
            # Если переменная не пуста, сохраняем ее в список переменных
            if current_variable.strip():  # Добавлена проверка на пустоту
                ww_vars.append(current_variable)
            # Создаем новую переменную и добавляем текущую строку с цифрой
            current_variable = line
        else:
            # Добавляем текущую строку с текстом к текущей переменной
            current_variable += line
    # Добавляем последнюю переменную в список (если она не пуста)
    if current_variable.strip():  # Добавлена проверка на пустоту
        ww_vars.append(current_variable)
# Выводим результат
for i, variable in enumerate(ww_vars, 1):
    print(f"ww_vars[{i-1}]:\n{variable}")





#-----СОЕДИНЕНИЕ ШАБЛОНА С РЕФЕРАТОМ
from docx import Document

# Открываем файлы
doc = Document('1.docx')
with open('2.txt', 'r') as file:
    text = file.read()
# Добавляем текст в документ
doc.add_paragraph(text)
# Сохраняем изменения
doc.save('1.docx')








#-----СОЕДИНЕНИЕ ШАБЛОНА С РЕФЕРАТОМ + Сохраняем документ с изменениями шрифта
import docx
from docx.shared import Pt
from docx import Document
import chardet
# Открываем файлы с явным указанием кодировки (в данном случае UTF-8)
with open('2.txt', 'rb') as file:
    result = chardet.detect(file.read())
# Открываем файлы с определенной кодировкой
with open('2.txt', 'r', encoding=result['encoding']) as file:
    text = file.read()
# Добавляем текст в документ
doc = Document('shablon.docx')
doc.add_paragraph(text)
# Сохраняем изменения
doc.save('shablonDOP.docx')
# Открываем документ с добавленным текстом
document = docx.Document('shablonDOP.docx')
# Изменяем стиль по умолчанию для всего документа
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
# Сохраняем документ с изменениями шрифта
document.save('REFERAT.docx')
print('Шрифт в файле успешно изменен на Times New Roman 14pt')







#--6.2-очистить-что-после--через переменнную

# Предположим, что ww_vars[17] - это строка данных
ww_vars_17 = ww_vars[17]
# Разделение строки на список строк
lines = ww_vars_17.split('\n')
# Ищем строку с числом 6.2
index_of_6_2 = None
for i, line in enumerate(lines):
    if "6.2" in line:
        index_of_6_2 = i
        break
# Если нашли строку с 6.2, ищем пустую строку после нее
if index_of_6_2 is not None:
    for i in range(index_of_6_2 + 1, len(lines)):
        if not lines[i].strip():
            # Нашли пустую строку, удаляем все строки после нее
            lines = lines[:i]
            break
# Собираем строки обратно в одну строку
result17 = '\n'.join(lines)
# Выводим результат
print(result17)







#=============================================================
#=============================================================
#=============БЛОК=ИСТОЧНИКОВ-ИНФЫ
import os
import re
import openai
import telebot
import time
import docx
from docx.shared import Pt
from docx import Document
import chardet



openai.api_key = ""

model_engine = "gpt-3.5-turbo"
with open('1.txt', 'w') as file:
    file.write('')
with open('2.txt', 'w') as file:
    file.write('')
text = "Приоритетные функции российского государства на пути выхода из кризиса"




# =====гл запрос======================================================
# =====гл запрос======================================================
# =====гл запрос======================================================
# =====гл запрос======================================================
ll2 = ""
prompt = "перечисли (назови) источники информации для моего реферата (" + text + ") , твой ответ должен содержать 5 электронных источника (статьи, сайты) и 5 печатных источника (книги, газеты, журналы), Источники пронумеруй так (ЭЛЕКТРОННЫЕ ИСТОЧНИКИ) 1 , 2, 3, 4, 5  (ПЕЧАТНЫЕ ИСТОЧНИКИ) 1, 2, 3, 4, 5 "
response = openai.ChatCompletion.create(
    model=model_engine,
    messages=[
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": prompt}
    ]
)
if response['choices']:
    ll2 = response['choices'][0]['message']['content']
print(ll2)
# ================СОЗДАНИЕ ВЫСШЕГО ВОПРОСА===========================================
# ================СОЗДАНИЕ ВЫСШЕГО ВОПРОСА===========================================
# ================СОЗДАНИЕ ВЫСШЕГО ВОПРОСА===========================================
# ================СОЗДАНИЕ ВЫСШЕГО ВОПРОСА===========================================











